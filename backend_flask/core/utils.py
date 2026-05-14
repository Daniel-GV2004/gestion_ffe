import os
import io
import copy
import datetime
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from bson.errors import InvalidId
from .constants import (
    etiquetas_centro, 
    etiquetas_alumno, 
    etiquetas_empresa,
    etiquetas_usuario,
    etiquetas_practica
)
from core.models import Practica, Alumno, Empresa

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PLANTILLAS_DIR = os.path.join(BASE_DIR, 'plantillas_word')

def realizar_sustitucion_profunda(elemento_xml, mapeo):
    for p_xml in elemento_xml.xpath('.//w:p'):
        nodos_texto = [nodo for nodo in p_xml.iter() if nodo.tag.endswith('}t') and nodo.text]
        if not nodos_texto:
            continue
            
        texto_completo = "".join(n.text for n in nodos_texto)
        cambiado = False
        
        for etiqueta, valor in mapeo.items():
            if etiqueta in texto_completo:
                texto_completo = texto_completo.replace(etiqueta, str(valor))
                cambiado = True
        
        if cambiado:
            nodos_texto[0].text = texto_completo
            for n in nodos_texto[1:]:
                n.text = ""

def procesar_plantilla_word(doc, reemplazos_simples=None, reemplazos_tablas=None):
    # --- 1. TABLAS DINÁMICAS (Búsqueda XML Extrema) ---
    if reemplazos_tablas:
        for etiqueta_clave, lista_reemplazos in reemplazos_tablas.items():
            if not lista_reemplazos:
                continue
            
            tabla_obj_xml = None
            fila_molde_xml = None

            for tbl in doc._element.xpath('.//w:tbl'):
                for tr in tbl.xpath('.//w:tr'):
                    textos_celdas = []
                    for tc in tr.xpath('.//w:tc'):
                        texto_tc = "".join([n.text for n in tc.iter() if n.tag.endswith('}t') and n.text])
                        textos_celdas.append(texto_tc)
                        
                    if any(etiqueta_clave in t for t in textos_celdas):
                        tabla_obj_xml = tbl
                        fila_molde_xml = tr
                        break
                if tabla_obj_xml is not None: 
                    break
            
            if tabla_obj_xml is None or fila_molde_xml is None:
                continue

            for _ in range(len(lista_reemplazos) - 1):
                nueva_fila = copy.deepcopy(fila_molde_xml)
                fila_molde_xml.addnext(nueva_fila)

            filas_finales = []
            for tr in tabla_obj_xml.xpath('.//w:tr'):
                textos_celdas = []
                for tc in tr.xpath('.//w:tc'):
                    texto_tc = "".join([n.text for n in tc.iter() if n.tag.endswith('}t') and n.text])
                    textos_celdas.append(texto_tc)
                if any(etiqueta_clave in t for t in textos_celdas):
                    filas_finales.append(tr)

            for i, reemplazos_fila in enumerate(lista_reemplazos):
                if i < len(filas_finales):
                    realizar_sustitucion_profunda(filas_finales[i], reemplazos_fila)

    # --- 2. REEMPLAZOS SIMPLES GLOBALES ---
    if reemplazos_simples:
        realizar_sustitucion_profunda(doc._element, reemplazos_simples)

        for section in doc.sections:
            for hp in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
                if hp and hasattr(hp, '_element'):
                    realizar_sustitucion_profunda(hp._element, reemplazos_simples)

    return doc

def generar_documento_en_memoria(datos_peticion, datos_completos_bd):
    doc_id = datos_peticion.get('docId')
    ruta_plantilla = os.path.join(PLANTILLAS_DIR, f"{doc_id}.docx")
    
    if not os.path.exists(ruta_plantilla):
        raise FileNotFoundError(f"No se ha encontrado la plantilla: {doc_id}.docx")

    doc = Document(ruta_plantilla)
    reemplazos_simples = etiquetas_centro()
    
    usuario_data = datos_peticion.get('usuario', {})
    reemplazos_simples.update(etiquetas_usuario(usuario_data))

    reemplazos_tablas = {"[!]": []}

    def procesar_elemento(elemento):
        if not elemento: return

        if isinstance(elemento, str):
            documento_db = None
            try:
                documento_db = Empresa.objects(id=elemento).first()
                if not documento_db:
                    documento_db = Alumno.objects(id=elemento).first()
            except InvalidId: pass

            if not documento_db:
                documento_db = Empresa.objects(cif=elemento).first()
            if not documento_db:
                documento_db = Alumno.objects(nif=elemento).first()

            if documento_db:
                elemento = documento_db.to_mongo().to_dict()
                if "nif" in elemento:
                    practica_db = Practica.objects(alumno=documento_db.id).first()
                    if practica_db:
                        elemento["practica_asociada"] = practica_db.to_mongo().to_dict()

        if isinstance(elemento, dict):
            if "cif" in elemento or "nombre_empresa" in elemento:
                reemplazos_simples.update(etiquetas_empresa(elemento))
                
            if "nif" in elemento or "apellidos" in elemento:
                etiquetas_al = etiquetas_alumno(elemento)
                if "practica_asociada" in elemento:
                    etiquetas_al.update(etiquetas_practica(elemento["practica_asociada"]))
                
                reemplazos_simples.update(etiquetas_al)
                
                etiquetas_al_tabla = etiquetas_al.copy()
                etiquetas_al_tabla["[!]"] = "" 
                reemplazos_tablas["[!]"].append(etiquetas_al_tabla)

            if "empresa" in elemento: procesar_elemento(elemento["empresa"])
            if "alumno" in elemento: procesar_elemento(elemento["alumno"])
            if "alumnos" in elemento and isinstance(elemento["alumnos"], list):
                for al in elemento["alumnos"]: procesar_elemento(al)

    if isinstance(datos_completos_bd, list):
        for item in datos_completos_bd: procesar_elemento(item)
    elif isinstance(datos_completos_bd, dict):
        procesar_elemento(datos_completos_bd)

    if not reemplazos_tablas.get("[!]"):
        if "[!]" in reemplazos_tablas: del reemplazos_tablas["[!]"]

    doc = procesar_plantilla_word(doc, reemplazos_simples, reemplazos_tablas)

    memoria = io.BytesIO()
    doc.save(memoria)
    memoria.seek(0) 
    return memoria

def procesar_anexo_oficial_junta(doc_id, practica_id):
    """
    Procesa Anexo II y Anexo IV buscando casillas por proximidad,
    alimentándose de los datos ya formateados en constants.py.
    """
    practica = Practica.objects.get(id=practica_id)
    alumno = Alumno.objects.get(id=practica.alumno.id)
    empresa = Empresa.objects.get(id=practica.empresa.id)
    
    # 1. Preparar diccionarios formateados mediante constants.py
    alumno_dict = alumno.to_mongo().to_dict()
    empresa_dict = empresa.to_mongo().to_dict()
    practica_dict = practica.to_mongo().to_dict()

    datos_centro = etiquetas_centro()
    datos_al = etiquetas_alumno(alumno_dict)
    datos_emp = etiquetas_empresa(empresa_dict)
    datos_prac = etiquetas_practica(practica_dict)

    # 2. Selección de plantilla
    if doc_id == "Anexo II Plan formativo":
        ruta_plantilla = os.path.join(PLANTILLAS_DIR, "Anexo_II_Original.docx")
    elif doc_id == "Anexo IV Informe valorativo":
        ruta_plantilla = os.path.join(PLANTILLAS_DIR, "Anexo_IV_Original.docx")
    else:
        raise ValueError("Documento oficial no soportado")

    doc = Document(ruta_plantilla)
    seccion_actual = ""

    # 3. Escaneo y rellenado posicional
    for table in doc.tables:
        if not table.rows: continue
        
        texto_cabecera = table.rows[0].cells[0].text.upper()
        
        # Determinar contexto de la sección
        if "PERSONA EN FORMACIÓN" in texto_cabecera or "ALUMNO" in texto_cabecera:
            seccion_actual = "alumno"
        elif "DATOS IDENTIFICATIVOS DE LA EMPRESA" in texto_cabecera:
            seccion_actual = "empresa"
        elif "DATOS IDENTIFICATIVOS DEL CENTRO" in texto_cabecera:
            seccion_actual = "centro"

        for row in table.rows:
            for i, cell in enumerate(row.cells):
                texto_celda = cell.text.strip().upper()
                
                # Si hay una celda a la derecha, es nuestra celda de destino
                if i + 1 < len(row.cells):
                    celda_destino = row.cells[i+1]
                    
                    if seccion_actual == "alumno":
                        if texto_celda == "NOMBRE": celda_destino.text = alumno_dict.get("nombre", "")
                        elif texto_celda == "APELLIDOS": celda_destino.text = alumno_dict.get("apellidos", "")
                        elif texto_celda in ["NIF", "DNI", "NIE"]: celda_destino.text = datos_al["[NIF]"]
                        elif texto_celda == "EMAIL": celda_destino.text = datos_al["[EMAIL]"]
                        elif texto_celda == "TELÉFONO": celda_destino.text = datos_al["[TELEFONO]"]
                        elif texto_celda == "NUSS": celda_destino.text = datos_al["[NUSS]"]
                        elif texto_celda == "CURSO": celda_destino.text = datos_prac["[CURSO_PRACTICA]"]
                            
                    elif seccion_actual == "empresa":
                        if texto_celda == "DENOMINACIÓN": celda_destino.text = datos_emp["[EMP_NOMBRE]"]
                        elif texto_celda == "CIF": celda_destino.text = datos_emp["[EMP_CIF]"]
                        elif texto_celda == "EMAIL": celda_destino.text = datos_emp["[EMP_EMAIL]"]
                        elif texto_celda == "TELÉFONO": celda_destino.text = datos_emp["[EMP_TELEFONO]"]
                        elif texto_celda == "NOMBRE Y APELLIDOS": # Tutor empresa
                            celda_destino.text = datos_emp["[TUTOR_NOMBRE]"]

                    elif seccion_actual == "centro":
                        if texto_celda == "CENTRO EDUCATIVO": celda_destino.text = datos_centro["[CENTRO_NOMBRE]"]
                        elif texto_celda == "CÓDIGO": celda_destino.text = datos_centro["[CENTRO_CODIGO]"]
                        elif texto_celda == "CIF": celda_destino.text = datos_centro["[CENTRO_NIF]"]
                        elif texto_celda == "EMAIL": celda_destino.text = datos_centro["[CENTRO_EMAIL]"]
                        elif texto_celda == "TELÉFONO": celda_destino.text = datos_centro["[CENTRO_TELEFONO]"]

                # Lógica de Checkboxes (reemplazo de caracteres)
                if "EN UNA EMPRESA" in texto_celda and "☐" in cell.text:
                    cell.text = cell.text.replace("☐", "☒")
                if "NO" in texto_celda and "exención" in cell.text.lower() and "☐" in cell.text:
                    cell.text = cell.text.replace("☐", "☒")

    # 4. Guardado final en memoria
    memoria = io.BytesIO()
    doc.save(memoria)
    memoria.seek(0)
    return memoria, alumno