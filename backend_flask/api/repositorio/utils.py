import os
import io
import copy
import datetime
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from bson.errors import InvalidId
from core.constants import (
    etiquetas_centro, 
    etiquetas_alumno, 
    etiquetas_empresa,
    etiquetas_usuario,
    etiquetas_practica
)
from core.models import Practica, Alumno, Empresa

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
PLANTILLAS_DIR = os.path.join(BASE_DIR, 'plantillas_word')

def realizar_sustitucion_profunda(elemento_xml, mapeo):
    for p_xml in elemento_xml.xpath('.//w:p'):
        nodos_texto = [nodo for nodo in p_xml.iter() if nodo.tag.endswith('}t') and nodo.text]
        if not nodos_texto:
            continue
            
        texto_completo = "".join(n.text for n in nodos_texto)
        cambiado = False
        
        for etiqueta, valor in mapeo.items():
            if etiqueta in texto_completo:
                texto_completo = texto_completo.replace(etiqueta, str(valor))
                cambiado = True
        
        if cambiado:
            nodos_texto[0].text = texto_completo
            for n in nodos_texto[1:]:
                n.text = ""

def procesar_plantilla_word(doc, reemplazos_simples=None, reemplazos_tablas=None):
    # --- 1. TABLAS DINÁMICAS (Búsqueda XML Extrema) ---
    if reemplazos_tablas:
        for etiqueta_clave, lista_reemplazos in reemplazos_tablas.items():
            if not lista_reemplazos:
                continue
            
            tabla_obj_xml = None
            fila_molde_xml = None

            for tbl in doc._element.xpath('.//w:tbl'):
                for tr in tbl.xpath('.//w:tr'):
                    textos_celdas = []
                    for tc in tr.xpath('.//w:tc'):
                        texto_tc = "".join([n.text for n in tc.iter() if n.tag.endswith('}t') and n.text])
                        textos_celdas.append(texto_tc)
                        
                    if any(etiqueta_clave in t for t in textos_celdas):
                        tabla_obj_xml = tbl
                        fila_molde_xml = tr
                        break
                if tabla_obj_xml is not None: 
                    break
            
            if tabla_obj_xml is None or fila_molde_xml is None:
                continue

            for _ in range(len(lista_reemplazos) - 1):
                nueva_fila = copy.deepcopy(fila_molde_xml)
                fila_molde_xml.addnext(nueva_fila)

            filas_finales = []
            for tr in tabla_obj_xml.xpath('.//w:tr'):
                textos_celdas = []
                for tc in tr.xpath('.//w:tc'):
                    texto_tc = "".join([n.text for n in tc.iter() if n.tag.endswith('}t') and n.text])
                    textos_celdas.append(texto_tc)
                if any(etiqueta_clave in t for t in textos_celdas):
                    filas_finales.append(tr)

            for i, reemplazos_fila in enumerate(lista_reemplazos):
                if i < len(filas_finales):
                    realizar_sustitucion_profunda(filas_finales[i], reemplazos_fila)

    # --- 2. REEMPLAZOS SIMPLES GLOBALES ---
    if reemplazos_simples:
        realizar_sustitucion_profunda(doc._element, reemplazos_simples)

        for section in doc.sections:
            for hp in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
                if hp and hasattr(hp, '_element'):
                    realizar_sustitucion_profunda(hp._element, reemplazos_simples)

    return doc

def generar_documento_en_memoria(datos_peticion, datos_completos_bd):
    doc_id = datos_peticion.get('docId')
    ruta_plantilla = os.path.join(PLANTILLAS_DIR, f"{doc_id}.docx")
    
    if not os.path.exists(ruta_plantilla):
        raise FileNotFoundError(f"No se ha encontrado la plantilla: {doc_id}.docx")

    doc = Document(ruta_plantilla)
    reemplazos_simples = etiquetas_centro()
    
    usuario_data = datos_peticion.get('usuario', {})
    reemplazos_simples.update(etiquetas_usuario(usuario_data))

    reemplazos_tablas = {"[!]": []}

    def procesar_elemento(elemento):
        if not elemento: return

        if isinstance(elemento, str):
            documento_db = None
            try:
                documento_db = Empresa.objects(id=elemento).first()
                if not documento_db:
                    documento_db = Alumno.objects(id=elemento).first()
            except InvalidId: pass

            if not documento_db:
                documento_db = Empresa.objects(cif=elemento).first()
            if not documento_db:
                documento_db = Alumno.objects(nif=elemento).first()

            if documento_db:
                elemento = documento_db.to_mongo().to_dict()
                if "nif" in elemento:
                    practica_db = Practica.objects(alumno=documento_db.id).first()
                    if practica_db:
                        elemento["practica_asociada"] = practica_db.to_mongo().to_dict()

        if isinstance(elemento, dict):
            if "cif" in elemento or "nombre_empresa" in elemento:
                reemplazos_simples.update(etiquetas_empresa(elemento))
                
            if "nif" in elemento or "apellidos" in elemento:
                etiquetas_al = etiquetas_alumno(elemento)
                if "practica_asociada" in elemento:
                    etiquetas_al.update(etiquetas_practica(elemento["practica_asociada"]))
                
                reemplazos_simples.update(etiquetas_al)
                
                etiquetas_al_tabla = etiquetas_al.copy()
                etiquetas_al_tabla["[!]"] = "" 
                reemplazos_tablas["[!]"].append(etiquetas_al_tabla)

            if "empresa" in elemento: procesar_elemento(elemento["empresa"])
            if "alumno" in elemento: procesar_elemento(elemento["alumno"])
            if "alumnos" in elemento and isinstance(elemento["alumnos"], list):
                for al in elemento["alumnos"]: procesar_elemento(al)

    if isinstance(datos_completos_bd, list):
        for item in datos_completos_bd: procesar_elemento(item)
    elif isinstance(datos_completos_bd, dict):
        procesar_elemento(datos_completos_bd)

    if not reemplazos_tablas.get("[!]"):
        if "[!]" in reemplazos_tablas: del reemplazos_tablas["[!]"]

    doc = procesar_plantilla_word(doc, reemplazos_simples, reemplazos_tablas)

    memoria = io.BytesIO()
    doc.save(memoria)
    memoria.seek(0) 
    return memoria

def procesar_anexo_oficial_junta(archivo_subido, practica_id):
    """
    Procesa Anexo II y Anexo IV buscando casillas por proximidad.
    Utiliza lectura directa del XML para evitar crasheos con celdas combinadas.
    """
    import io
    from docx import Document
    from docx.table import _Cell
    from core.models import Practica, Alumno, Empresa
    # from .constants import etiquetas_centro, etiquetas_alumno, etiquetas_empresa, etiquetas_practica

    practica = Practica.objects.get(id=practica_id)
    alumno = Alumno.objects.get(id=practica.alumno.id)
    empresa = Empresa.objects.get(id=practica.empresa.id)
    
    # 1. Preparar diccionarios formateados mediante constants.py
    alumno_dict = alumno.to_mongo().to_dict()
    empresa_dict = empresa.to_mongo().to_dict()
    practica_dict = practica.to_mongo().to_dict()

    datos_centro = etiquetas_centro()
    datos_al = etiquetas_alumno(alumno_dict)
    datos_emp = etiquetas_empresa(empresa_dict)
    datos_prac = etiquetas_practica(practica_dict)

    doc = Document(archivo_subido)
    seccion_actual = ""

    # 2. Escaneo a través del código fuente XML de las tablas
    for table in doc.tables:
        tbl_xml = table._tbl
        
        filas_xml = tbl_xml.xpath('.//w:tr')
        if not filas_xml: continue
        
        primeras_celdas = filas_xml[0].xpath('.//w:tc')
        if not primeras_celdas: continue
        
        celda_cabecera = _Cell(primeras_celdas[0], table)
        texto_cabecera = celda_cabecera.text.upper()
        
        # Determinar contexto general de la tabla
        if "PERSONA EN FORMACIÓN" in texto_cabecera or "ALUMNO" in texto_cabecera:
            seccion_actual = "alumno"
        elif "DATOS IDENTIFICATIVOS DE LA EMPRESA" in texto_cabecera or "DATOS DE LA EMPRESA" in texto_cabecera:
            seccion_actual = "empresa"
        elif "DATOS IDENTIFICATIVOS DEL CENTRO" in texto_cabecera or "DATOS DEL CENTRO" in texto_cabecera:
            seccion_actual = "centro"

        # 3. Rellenado celda por celda
        for tr in filas_xml:
            celdas_xml = tr.xpath('.//w:tc')
            
            for i, tc in enumerate(celdas_xml):
                cell = _Cell(tc, table) 
                texto_celda = cell.text.strip().upper()
                
                # --- AQUÍ ESTÁ LA MAGIA: DETECTAR SUB-SECCIONES ---
                # Si en cualquier celda leemos que empieza la parte del tutor, cambiamos la sección
                if "TUTOR/A DE LA EMPRESA" in texto_celda:
                    seccion_actual = "tutor_empresa"
                elif "TUTOR/A DEL CENTRO" in texto_celda:
                    seccion_actual = "tutor_centro"
                
                # Si hay una celda a la derecha, la preparamos como destino
                if i + 1 < len(celdas_xml):
                    celda_destino = _Cell(celdas_xml[i+1], table)
                    
                    if seccion_actual == "alumno":
                        if texto_celda == "NOMBRE": celda_destino.text = alumno_dict.get("nombre", "")
                        elif texto_celda == "APELLIDOS": celda_destino.text = alumno_dict.get("apellidos", "")
                        elif texto_celda in ["NIF", "DNI", "NIE"]: celda_destino.text = datos_al.get("[NIF]", "")
                        elif texto_celda == "EMAIL": celda_destino.text = datos_al.get("[EMAIL]", "")
                        elif texto_celda == "TELÉFONO": celda_destino.text = datos_al.get("[TELEFONO]", "")
                        elif texto_celda == "NUSS": celda_destino.text = datos_al.get("[NUSS]", "")
                        elif texto_celda == "CURSO": celda_destino.text = datos_prac.get("[CURSO_PRACTICA]", "")
                            
                    elif seccion_actual == "empresa":
                        if texto_celda in ["DENOMINACIÓN", "EMPRESA U ORGANISMO EQUIPARADO"]: 
                            celda_destino.text = datos_emp.get("[EMP_NOMBRE]", "")
                        elif texto_celda == "CIF": 
                            celda_destino.text = datos_emp.get("[EMP_CIF]", "")
                        elif texto_celda == "EMAIL": 
                            celda_destino.text = datos_emp.get("[EMP_EMAIL]", "")
                        elif texto_celda == "TELÉFONO": 
                            celda_destino.text = datos_emp.get("[EMP_TELEFONO]", "")

                    # --- NUEVA SECCIÓN EXCLUSIVA PARA EL TUTOR ---
                    elif seccion_actual == "tutor_empresa":
                        if texto_celda == "NOMBRE Y APELLIDOS": 
                            celda_destino.text = datos_emp.get("[TUTOR_NOMBRE]", "")
                        elif texto_celda == "EMAIL": 
                            celda_destino.text = datos_emp.get("[TUTOR_EMAIL]", "") # Ahora sí, el correo correcto

                    elif seccion_actual == "centro":
                        if texto_celda == "CENTRO EDUCATIVO": celda_destino.text = datos_centro.get("[CENTRO_NOMBRE]", "")
                        elif texto_celda == "CÓDIGO": celda_destino.text = datos_centro.get("[CENTRO_CODIGO]", "")
                        elif texto_celda == "CIF": celda_destino.text = datos_centro.get("[CENTRO_NIF]", "")
                        elif texto_celda == "EMAIL": celda_destino.text = datos_centro.get("[CENTRO_EMAIL]", "")
                        elif texto_celda == "TELÉFONO": celda_destino.text = datos_centro.get("[CENTRO_TELEFONO]", "")

                # Lógica de Checkboxes (reemplazo de caracteres en la celda actual)
                if "EN UNA EMPRESA" in texto_celda and "☐" in cell.text:
                    cell.text = cell.text.replace("☐", "☒")
                if "NO" in texto_celda and "exención" in cell.text.lower() and "☐" in cell.text:
                    cell.text = cell.text.replace("☐", "☒")

    # 4. Guardado final en memoria RAM
    memoria = io.BytesIO()
    doc.save(memoria)
    memoria.seek(0)
    
    return memoria, alumno